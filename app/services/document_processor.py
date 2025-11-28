import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from fastapi import HTTPException
from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
            '.csv': self._extract_csv,
            '.txt': self._extract_txt
        }

    def is_supported_format(self, filename: str) -> bool:
        """Verifica si el formato del archivo es soportado"""
        file_ext = os.path.splitext(filename)[1].lower()
        return file_ext in self.supported_formats

    async def extract_text(self, file_path: str) -> str:
        """Extrae texto de diferentes tipos de archivos"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if not self.is_supported_format(file_path):
            raise HTTPException(status_code=400, detail=f"Formato no soportado: {file_ext}")
        
        try:
            # Las funciones de extracción son síncronas, no necesitan await
            extract_function = self.supported_formats[file_ext]
            return extract_function(file_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error procesando archivo: {str(e)}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extrae texto de archivos PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                print(f"📄 PDF tiene {len(reader.pages)} páginas")
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"--- Página {page_num + 1} ---\n"
                        text += page_text + "\n\n"
                
                if not text.strip():
                    text = "⚠️ No se pudo extraer texto del PDF. Puede ser un PDF escaneado o con imágenes."
                
                return text.strip()
        except Exception as e:
            raise Exception(f"Error leyendo PDF: {str(e)}")

    def _extract_docx(self, file_path: str) -> str:
        """Extrae texto de archivos Word (.docx, .doc)"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                text += "\n--- TABLA ---\n"
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text += cell.text + " | "
                    if row_text:
                        text += row_text.rstrip(" | ") + "\n"
                text += "--- FIN TABLA ---\n"
            
            if not text.strip():
                text = "⚠️ El documento Word parece estar vacío o no contiene texto extraíble."
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error leyendo documento Word: {str(e)}")

    def _extract_excel(self, file_path: str) -> str:
        """Extrae texto de archivos Excel (.xlsx, .xls)"""
        try:
            text = ""
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            print(f"📊 Excel tiene {len(excel_file.sheet_names)} hojas: {excel_file.sheet_names}")
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\n--- HOJA: {sheet_name} ---\n"
                text += f"Dimensiones: {df.shape[0]} filas x {df.shape[1]} columnas\n\n"
                
                # Mostrar primeras filas para el resumen
                if not df.empty:
                    # Encabezados
                    headers = " | ".join(str(col) for col in df.columns)
                    text += f"Encabezados: {headers}\n\n"
                    
                    # Primeras 5 filas como muestra
                    for i, row in df.head(5).iterrows():
                        row_data = " | ".join(str(val) if pd.notna(val) else "N/A" for val in row)
                        text += f"Fila {i+1}: {row_data}\n"
                    
                    if len(df) > 5:
                        text += f"... y {len(df) - 5} filas más\n"
                
                text += "\n"
            
            return text.strip() if text.strip() else "⚠️ El archivo Excel está vacío o no contiene datos."
        except Exception as e:
            raise Exception(f"Error leyendo archivo Excel: {str(e)}")

    def _extract_csv(self, file_path: str) -> str:
        """Extrae texto de archivos CSV"""
        try:
            df = pd.read_csv(file_path)
            text = f"Archivo CSV: {len(df)} filas x {len(df.columns)} columnas\n\n"
            
            if not df.empty:
                # Encabezados
                headers = " | ".join(str(col) for col in df.columns)
                text += f"Encabezados: {headers}\n\n"
                
                # Primeras filas como muestra
                for i, row in df.head(10).iterrows():
                    row_data = " | ".join(str(val) if pd.notna(val) else "N/A" for val in row)
                    text += f"Fila {i+1}: {row_data}\n"
                
                if len(df) > 10:
                    text += f"... y {len(df) - 10} filas más\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error leyendo archivo CSV: {str(e)}")

    def _extract_txt(self, file_path: str) -> str:
        """Extrae texto de archivos de texto plano"""
        try:
            # Probar diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            print(f"✅ Texto leído con codificación: {encoding}")
                            return content.strip()
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # Si ninguna codificación funciona, usar binary
            with open(file_path, 'rb') as file:
                content = file.read()
                # Intentar decodificar con UTF-8 ignorando errores
                return content.decode('utf-8', errors='ignore').strip()
                
        except Exception as e:
            raise Exception(f"Error leyendo archivo de texto: {str(e)}")

    def classify_document(self, text: str, filename: str) -> str:
        """Clasifica el documento como científico o general basado en contenido y nombre"""
        scientific_keywords = [
            'abstract', 'introduction', 'methodology', 'results', 'conclusion',
            'research', 'study', 'experiment', 'hypothesis', 'literature review',
            'citation', 'references', 'method', 'data analysis', 'findings',
            'objective', 'background', 'discussion', 'limitations', 'future work',
            'thesis', 'dissertation', 'paper', 'journal', 'conference'
        ]
        
        filename_lower = filename.lower()
        text_lower = text.lower()
        
        # Verificar palabras clave en el nombre del archivo
        scientific_filename_indicators = [
            'paper', 'research', 'study', 'thesis', 'dissertation', 'journal',
            'conference', 'experiment', 'analysis', 'report', 'cientific'
        ]
        
        if any(indicator in filename_lower for indicator in scientific_filename_indicators):
            return "scientific"
        
        # Verificar palabras clave en el contenido
        scientific_content_score = sum(1 for keyword in scientific_keywords if keyword in text_lower)
        
        # Umbral más bajo para documentos más cortos
        threshold = 2 if len(text_lower) < 1000 else 3
        
        if scientific_content_score >= threshold:
            return "scientific"
        else:
            return "general"